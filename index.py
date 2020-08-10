from docx import Document
from docx.shared import Inches, Cm, Length
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import sys
from datetime import datetime

def create_document(parent_path):
    document = Document()
    locale = os.path.basename(parent_path)
    style = document.styles['Normal']
    font = style.font
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    #Font Styling
    font = document.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(9)

    #Capture Local Time
    date = datetime.utcnow()
    currentTime = date.strftime("%a %b %d %Y %X")

    #Header section
    header_section = document.sections[0]
    header = header_section.header
    # header.top_margin = Inches(0.5)
    # header.bottom_margin = Cm(0.54)

    #Footer section
    footer_section = document.sections[0]
    footer = footer_section.footer


    # Header Text
    header_text = header.add_paragraph()
    header_text.text = (currentTime, "GMT+0000 (Coordinated Universal Time)")
    header_text.alignment = 2
    run = header_text.add_run()
    run.add_break()
    run.add_break()

    #Footer text
    footer_text = footer.paragraphs[0]
    footer_text.text = locale

    return document

def cover_page(document, parent_path, login, study_detais,device_name, environment_name):

    font = document.styles['Normal'].font
    locale = os.path.basename(parent_path)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.75
    paragraph_format.left_indent = Inches(2.0)
    run = paragraph.add_run()
    run.add_break()
    run.add_break()
    run.add_break()
    font.size = Pt(10)
    paragraph.add_run('Login Name: ' + '' + login)
    paragraph.add_run('\n')
    paragraph.add_run('Study Name: ' + '' + study_detais)
    paragraph.add_run('\n')
    paragraph.add_run('Supported Language: ' + '' + locale )
    paragraph.add_run('\n')
    paragraph.add_run('Environment: ' + '' + environment_name)
    paragraph.add_run('\n')
    paragraph.add_run('Device: ' + '' + device_name)
    paragraph.add_run('\n')
    paragraph.add_run('App Version: ')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_page_break()

def paragraph_header(document, folderName, filename):
    paragraph = document.add_paragraph()
    paragraph.add_run("Folder Name:  " + folderName)
    paragraph.add_run('\n')
    paragraph.add_run("File Name:  " + filename)
    paragraph.add_run('\n')
    paragraph.add_run('\n')




def explore_study(parent_path, login, study_detais,device_name, environment_name):

    study_path = os.path.dirname(parent_path)
    study_name = os.path.basename(study_path)
    output_doc_name = ' - '.join([study_name, os.path.split(parent_path)[1]])

    document = create_document(parent_path)

    # Cover Page - Displays the Config details received from the WebApp tool
    cover_page(document, parent_path, login, study_detais,device_name, environment_name)


    for root, subdir, files in os.walk(parent_path):
        if files:

            for filename in sorted(files): # Add's Images to Document

                folderName = os.path.basename(root)
                filepath = os.path.join(root, filename)
                if os.path.splitext(filepath)[1] in (".jpeg", ".png"):
                    paragraph_header(document,folderName,filename)
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run()
                    run.add_break()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run.add_picture(filepath, height=Inches(7.0),width=Inches(4.0))  # Add Image
                    run.add_break()
                    run.add_break()
                    run.add_break()
    document.save(f"{output_doc_name}.docx")


def process_path(parent_path, login_name, study, device, environment):

    for doc in os.listdir(parent_path): # this is the study folder
        if not doc.startswith('.'):
            fqp = os.path.join(os.getcwd(),parent_path, doc)
            login = login_name
            study_detais = study
            device_name = device
            environment_name = environment
            explore_study(fqp, login, study_detais,device_name, environment_name )
            create_document(fqp)

# if __name__ == "__main__":
#     process_path(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4], sys.argv[5])

process_path("example/Images", 'pvoola@mdsol.com', 'kms_automation_21012020', 'iPhone 8 Plus', 'sandbox')